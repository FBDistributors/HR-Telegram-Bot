import asyncio
import logging
import io
import docx
import google.generativeai as genai
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters.command import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove
import os
from dotenv import load_dotenv

load_dotenv()

# --- SOZLAMALAR ---
BOT_TOKEN = os.getenv("BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
HR_GROUP_ID = os.getenv("HR_GROUP_ID")
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# --- TILLAR UCHUN LUG'AT (TO'LIQ VERSIYASI) ---
texts = {
    'uz': {
        'welcome': "Assalomu alaykum! Muloqot uchun qulay tilni tanlang.",
        'ask_name': "To'liq ism-sharifingizni kiriting (masalan, Olimov Salim).",
        'ask_has_resume': "Rahmat. Arizani davom ettirish uchun tayyor rezyumeingiz mavjudmi?",
        'button_yes_resume': "‚úÖ Ha, rezyume yuborish",
        'button_no_resume': "‚ùå Yo'q, suhbatdan o'tish",
        'prompt_for_resume': "Marhamat, rezyumeni PDF yoki DOCX formatida yuboring.",
        'start_convo_application': "Hechqisi yo'q! Keling, o'rniga bir nechta savollar orqali siz haqingizda ma'lumot olamiz.",
        'ask_vacancy': "Murojaat qilayotgan vakansiya nomini kiriting (masalan, Buxgalter).",
        'ask_experience': "Ish tajribangiz haqida yozing (oxirgi ish joyingiz, lavozimingiz, necha yil ishlaganingiz).",
        'ask_salary': "Oylik maosh bo'yicha kutilmalaringizni kiriting (so'mda, raqam yoki matn bilan yozing)",
        'ask_location': "Yashash manzilingizni kiriting (shahar, tuman).",
        'ask_skills': "Lavozimga oid eng muhim ko'nikmalaringizni sanab o'ting (masalan, Excel, 1C, mijozlar bilan ishlash).",
        'ask_availability': "Yaqin kunlarda ish boshlashga tayyormisiz?",
        'button_yes': "‚úÖ Ha",
        'button_no': "‚ùå Yo'q",
        'ask_contact': "Siz bilan bog'lanish uchun, quyidagi tugma orqali telefon raqamingizni yuboring:",
        'button_share_contact': "üì± Kontaktimni ulashish",
        'goodbye_user': "Barcha ma'lumotlaringiz uchun rahmat! Arizangiz muvaffaqiyatli qabul qilindi. Nomzodingiz ma'qul topilsa, biz siz bilan tez orada bog'lanamiz. ‚úÖ",
        'analyzing': "Ma'lumotlar qabul qilindi. Hozir sun'iy intellekt yordamida tahlil qilinmoqda, bir oz kuting...",
        'file_error': "Iltimos, rezyumeni faqat PDF yoki DOCX formatida yuboring.",
        'hr_notification_file': """üîî **Yangi nomzod (Rezyume bilan)!**

üë§ **FIO:** {name}
üìÑ **Rezyume:** Fayl biriktirildi.
-------------------
{summary}""",
        'hr_notification_convo': """üîî **Yangi nomzod (Suhbat orqali)!**

üë§ **FIO:** {name}
üë®‚Äçüíº **Vakansiya:** {vacancy}
-------------------
**Nomzod javoblari:**
- **Tajribasi:** {experience}
- **Maosh kutilmasi:** {salary}
- **Manzili:** {location}
- **Ko'nikmalari:** {skills}
- **Ishga tayyorligi:** {availability}
- **Aloqa:** {contact}
-------------------
{summary}""",
        'gemini_file_prompt': """Sen tajribali HR-menejersan. Ilova qilingan fayl nomzodning rezyumesi hisoblanadi. 
        Ushbu rezyumeni o'qib chiqib, nomzod haqida o'zbek tilida, lotin alifbosida qisqacha va aniq xulosa yoz.
        Tahlil quyidagi formatda bo'lsin, sarlavhalar va ro'yxatlar uchun emoji'lardan foydalan:
        ü§ñ **Umumiy xulosa:** [Nomzodning tajribasi, ko'nikmalari va ma'lumotlari asosida 2-3 gaplik xulosa]
        ‚ú® **Kuchli tomonlari:**
        ‚úÖ [Rezyumedan topilgan birinchi kuchli jihat]
        ‚úÖ [Rezyumedan topilgan ikkinchi kuchli jihat]
        üìä **Dastlabki baho:** [Mos keladi / O'ylab ko'rish kerak / Tajribasi kam]""",
        'gemini_text_prompt': """Sen tajribali HR-menejersan. Quyida nomzodning rezyumesidan olingan matn keltirilgan. 
        Ushbu matnni tahlil qilib, nomzod haqida o'zbek tilida, lotin alifbosida qisqacha va aniq xulosa yoz.
        Tahlil quyidagi formatda bo'lsin, sarlavhalar va ro'yxatlar uchun emoji'lardan foydalan:
        ü§ñ **Umumiy xulosa:** [Nomzodning tajribasi, ko'nikmalari va ma'lumotlari asosida 2-3 gaplik xulosa]
        ‚ú® **Kuchli tomonlari:**
        ‚úÖ [Rezyumedan topilgan birinchi kuchli jihat]
        ‚úÖ [Rezyumedan topilgan ikkinchi kuchli jihat]
        üìä **Dastlabki baho:** [Mos keladi / O'ylab ko'rish kerak / Tajribasi kam]
        
        Rezyume matni:
        {resume_text}
        """,
        'gemini_convo_prompt': """Sen tajribali HR-menejersan. Quyida nomzodning suhbat orqali bergan javoblari keltirilgan. 
        Ushbu ma'lumotlarni tahlil qilib, nomzod haqida o'zbek tilida, lotin alifbosida qisqacha va aniq xulosa yoz.
        Tahlil quyidagi formatda bo'lsin, emoji'lardan foydalan:
        ü§ñ **Umumiy xulosa:** [Nomzodning javoblari va vakansiyaga mosligi asosida 2-3 gaplik xulosa]
        ‚ú® **Kuchli tomonlari:**
        ‚úÖ [Suhbatdan topilgan birinchi kuchli jihat]
        ‚úÖ [Suhbatdan topilgan ikkinchi kuchli jihat]
        üìä **Dastlabki baho:** [Mos keladi / O'ylab ko'rish kerak / Tajribasi kam]"""
    },
    'ru': {
        'welcome': "–ó–¥—Ä–∞–≤—Å—Ç–≤—É–π—Ç–µ! –ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –≤—ã–±–µ—Ä–∏—Ç–µ —É–¥–æ–±–Ω—ã–π —è–∑—ã–∫ –¥–ª—è –æ–±—â–µ–Ω–∏—è.",
        'ask_name': "–í–≤–µ–¥–∏—Ç–µ –í–∞—à–∏ –ø–æ–ª–Ω—ã–µ –∏–º—è –∏ —Ñ–∞–º–∏–ª–∏—é (–Ω–∞–ø—Ä–∏–º–µ—Ä, –°–∞–ª–∏–º–æ–≤ –û–ª–∏–º).",
        'ask_has_resume': "–û—Ç–ª–∏—á–Ω–æ! –ß—Ç–æ–±—ã –ø—Ä–æ–¥–æ–ª–∂–∏—Ç—å, –∏–º–µ–µ—Ç—Å—è –ª–∏ —É –≤–∞—Å –≥–æ—Ç–æ–≤–æ–µ —Ä–µ–∑—é–º–µ?",
        'button_yes_resume': "‚úÖ –î–∞, –æ—Ç–ø—Ä–∞–≤–∏—Ç—å —Ä–µ–∑—é–º–µ",
        'button_no_resume': "‚ùå –ù–µ—Ç, –ø—Ä–æ–π—Ç–∏ —Å–æ–±–µ—Å–µ–¥–æ–≤–∞–Ω–∏–µ",
        'prompt_for_resume': "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ –≤–∞—à–µ —Ä–µ–∑—é–º–µ –≤ —Ñ–æ—Ä–º–∞—Ç–µ PDF –∏–ª–∏ DOCX.",
        'start_convo_application': "–ù–∏—á–µ–≥–æ —Å—Ç—Ä–∞—à–Ω–æ–≥–æ! –í–º–µ—Å—Ç–æ —ç—Ç–æ–≥–æ, –¥–∞–≤–∞–π—Ç–µ –ø–æ–ª—É—á–∏–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –≤–∞—Å —á–µ—Ä–µ–∑ –Ω–µ—Å–∫–æ–ª—å–∫–æ –≤–æ–ø—Ä–æ—Å–æ–≤.",
        'ask_vacancy': "–í–≤–µ–¥–∏—Ç–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –≤–∞–∫–∞–Ω—Å–∏–∏, –Ω–∞ –∫–æ—Ç–æ—Ä—É—é –í—ã –ø—Ä–µ—Ç–µ–Ω–¥—É–µ—Ç–µ (–Ω–∞–ø—Ä–∏–º–µ—Ä, '–ú–µ–Ω–µ–¥–∂–µ—Ä –ø–æ –ø—Ä–æ–¥–∞–∂–∞–º').",
        'ask_experience': "–û–ø–∏—à–∏—Ç–µ –≤–∞—à –æ–ø—ã—Ç —Ä–∞–±–æ—Ç—ã (–ø–æ—Å–ª–µ–¥–Ω–µ–µ –º–µ—Å—Ç–æ —Ä–∞–±–æ—Ç—ã, –¥–æ–ª–∂–Ω–æ—Å—Ç—å, —Å–∫–æ–ª—å–∫–æ –ª–µ—Ç —Ä–∞–±–æ—Ç–∞–ª–∏).",
        'ask_salary': "–£–∫–∞–∂–∏—Ç–µ –≤–∞—à–∏ –æ–∂–∏–¥–∞–Ω–∏—è –ø–æ –∑–∞—Ä–∞–±–æ—Ç–Ω–æ–π –ø–ª–∞—Ç–µ? (–≤ —Å—É–º–∞—Ö, –Ω–∞–ø–∏—à–∏—Ç–µ —Ü–∏—Ñ—Ä–æ–π –∏–ª–∏ —Ç–µ–∫—Å—Ç–æ–º)",
        'ask_location': "–í–≤–µ–¥–∏—Ç–µ –≤–∞—à –∞–¥—Ä–µ—Å –ø—Ä–æ–∂–∏–≤–∞–Ω–∏—è (–≥–æ—Ä–æ–¥, —Ä–∞–π–æ–Ω).",
        'ask_skills': "–ü–µ—Ä–µ—á–∏—Å–ª–∏—Ç–µ –≤–∞—à–∏ –∫–ª—é—á–µ–≤—ã–µ –Ω–∞–≤—ã–∫–∏ –¥–ª—è –¥–∞–Ω–Ω–æ–π –ø–æ–∑–∏—Ü–∏–∏ (–Ω–∞–ø—Ä–∏–º–µ—Ä, Excel, 1C, —Ä–∞–±–æ—Ç–∞ —Å –∫–ª–∏–µ–Ω—Ç–∞–º–∏).",
        'ask_availability': "–ì–æ—Ç–æ–≤—ã –ª–∏ –≤—ã –ø—Ä–∏—Å—Ç—É–ø–∏—Ç—å –∫ —Ä–∞–±–æ—Ç–µ –≤ –±–ª–∏–∂–∞–π—à–µ–µ –≤—Ä–µ–º—è?",
        'button_yes': "‚úÖ –î–∞",
        'button_no': "‚ùå –ù–µ—Ç",
        'ask_contact': "–î–ª—è —Å–≤—è–∑–∏, –ø–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ –≤–∞—à –Ω–æ–º–µ—Ä —Ç–µ–ª–µ—Ñ–æ–Ω–∞ —Å –ø–æ–º–æ—â—å—é –∫–Ω–æ–ø–∫–∏ –Ω–∏–∂–µ:",
        'button_share_contact': "üì± –ü–æ–¥–µ–ª–∏—Ç—å—Å—è –º–æ–∏–º –∫–æ–Ω—Ç–∞–∫—Ç–æ–º",
        'goodbye_user': "–°–ø–∞—Å–∏–±–æ –∑–∞ –≤—Å–µ –¥–∞–Ω–Ω—ã–µ! –í–∞—à–∞ –∑–∞—è–≤–∫–∞ —É—Å–ø–µ—à–Ω–æ –ø—Ä–∏–Ω—è—Ç–∞. –ú—ã —Å–≤—è–∂–µ–º—Å—è —Å –≤–∞–º–∏ –≤ –±–ª–∏–∂–∞–π—à–µ–µ –≤—Ä–µ–º—è, –µ—Å–ª–∏ –≤–∞—à–∞ –∫–∞–Ω–¥–∏–¥–∞—Ç—É—Ä–∞ –±—É–¥–µ—Ç –æ–¥–æ–±—Ä–µ–Ω–∞. ‚úÖ",
        'analyzing': "–î–∞–Ω–Ω—ã–µ –ø–æ–ª—É—á–µ–Ω—ã. –°–µ–π—á–∞—Å –æ–Ω–∏ –∞–Ω–∞–ª–∏–∑–∏—Ä—É—é—Ç—Å—è —Å –ø–æ–º–æ—â—å—é –∏—Å–∫—É—Å—Å—Ç–≤–µ–Ω–Ω–æ–≥–æ –∏–Ω—Ç–µ–ª–ª–µ–∫—Ç–∞, –ø–æ–¥–æ–∂–¥–∏—Ç–µ –Ω–µ–º–Ω–æ–≥–æ...",
        'file_error': "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ä–µ–∑—é–º–µ —Ç–æ–ª—å–∫–æ –≤ —Ñ–æ—Ä–º–∞—Ç–µ PDF –∏–ª–∏ DOCX.",
        'hr_notification_file': """üîî **–ù–æ–≤—ã–π –∫–∞–Ω–¥–∏–¥–∞—Ç (—Å —Ä–µ–∑—é–º–µ)!**

üë§ **–§–ò–û:** {name}
üìÑ **–†–µ–∑—é–º–µ:** –§–∞–π–ª –ø—Ä–∏–∫—Ä–µ–ø–ª–µ–Ω.
-------------------
{summary}""",
        'hr_notification_convo': """üîî **–ù–æ–≤—ã–π –∫–∞–Ω–¥–∏–¥–∞—Ç (—á–µ—Ä–µ–∑ —á–∞—Ç)!**

üë§ **–§–ò–û:** {name}
üë®‚Äçüíº **–í–∞–∫–∞–Ω—Å–∏—è:** {vacancy}
-------------------
**–û—Ç–≤–µ—Ç—ã –∫–∞–Ω–¥–∏–¥–∞—Ç–∞:**
- **–û–ø—ã—Ç:** {experience}
- **–û–∂–∏–¥–∞–µ–º–∞—è –∑–∞—Ä–ø–ª–∞—Ç–∞:** {salary}
- **–ê–¥—Ä–µ—Å:** {location}
- **–ù–∞–≤—ã–∫–∏:** {skills}
- **–ì–æ—Ç–æ–≤–Ω–æ—Å—Ç—å –∫ —Ä–∞–±–æ—Ç–µ:** {availability}
- **–ö–æ–Ω—Ç–∞–∫—Ç:** {contact}
-------------------
{summary}""",
        'gemini_file_prompt': """–¢—ã –æ–ø—ã—Ç–Ω—ã–π HR-–º–µ–Ω–µ–¥–∂–µ—Ä. –ü—Ä–∏–ª–æ–∂–µ–Ω–Ω—ã–π PDF-—Ñ–∞–π–ª —è–≤–ª—è–µ—Ç—Å—è —Ä–µ–∑—é–º–µ –∫–∞–Ω–¥–∏–¥–∞—Ç–∞. 
        –ü—Ä–æ—á–∏—Ç–∞–π —ç—Ç–æ —Ä–µ–∑—é–º–µ –∏ –Ω–∞–ø–∏—à–∏ –∫—Ä–∞—Ç–∫–æ–µ –∏ —á–µ—Ç–∫–æ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ –æ –∫–∞–Ω–¥–∏–¥–∞—Ç–µ –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ.
        –ê–Ω–∞–ª–∏–∑ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –≤ —Å–ª–µ–¥—É—é—â–µ–º —Ñ–æ—Ä–º–∞—Ç–µ, –∏—Å–ø–æ–ª—å–∑—É–π —ç–º–æ–¥–∑–∏ –¥–ª—è –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤ –∏ —Å–ø–∏—Å–∫–æ–≤:
        ü§ñ **–û–±—â–µ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ:** [–ó–∞–∫–ª—é—á–µ–Ω–∏–µ –∏–∑ 2-3 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –Ω–∞ –æ—Å–Ω–æ–≤–µ –æ–ø—ã—Ç–∞, –Ω–∞–≤—ã–∫–æ–≤ –∏ –æ–±—Ä–∞–∑–æ–≤–∞–Ω–∏—è –∫–∞–Ω–¥–∏–¥–∞—Ç–∞]
        ‚ú® **–°–∏–ª—å–Ω—ã–µ —Å—Ç–æ—Ä–æ–Ω—ã:**
        ‚úÖ [–ü–µ—Ä–≤–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ —Ä–µ–∑—é–º–µ]
        ‚úÖ [–í—Ç–æ—Ä–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ —Ä–µ–∑—é–º–µ]
        üìä **–ü—Ä–µ–¥–≤–∞—Ä–∏—Ç–µ–ª—å–Ω–∞—è –æ—Ü–µ–Ω–∫–∞:** [–ü–æ–¥—Ö–æ–¥–∏—Ç / –°—Ç–æ–∏—Ç —Ä–∞—Å—Å–º–æ—Ç—Ä–µ—Ç—å / –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –æ–ø—ã—Ç–∞]""",
        'gemini_text_prompt': """–¢—ã –æ–ø—ã—Ç–Ω—ã–π HR-–º–µ–Ω–µ–¥–∂–µ—Ä. –ù–∏–∂–µ –ø—Ä–∏–≤–µ–¥–µ–Ω —Ç–µ–∫—Å—Ç –∏–∑ —Ä–µ–∑—é–º–µ –∫–∞–Ω–¥–∏–¥–∞—Ç–∞. 
        –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —ç—Ç–æ—Ç —Ç–µ–∫—Å—Ç –∏ –Ω–∞–ø–∏—à–∏ –∫—Ä–∞—Ç–∫–æ–µ –∏ —á–µ—Ç–∫–æ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ –æ –∫–∞–Ω–¥–∏–¥–∞—Ç–µ –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ.
        –ê–Ω–∞–ª–∏–∑ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –≤ —Å–ª–µ–¥—É—é—â–µ–º —Ñ–æ—Ä–º–∞—Ç–µ, –∏—Å–ø–æ–ª—å–∑—É–π —ç–º–æ–¥–∑–∏ –¥–ª—è –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤ –∏ —Å–ø–∏—Å–∫–æ–≤:
        ü§ñ **–û–±—â–µ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ:** [–ó–∞–∫–ª—é—á–µ–Ω–∏–µ –∏–∑ 2-3 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –Ω–∞ –æ—Å–Ω–æ–≤–µ –æ–ø—ã—Ç–∞, –Ω–∞–≤—ã–∫–æ–≤ –∏ –æ–±—Ä–∞–∑–æ–≤–∞–Ω–∏—è –∫–∞–Ω–¥–∏–¥–∞—Ç–∞]
        ‚ú® **–°–∏–ª—å–Ω—ã–µ —Å—Ç–æ—Ä–æ–Ω—ã:**
        ‚úÖ [–ü–µ—Ä–≤–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ —Ä–µ–∑—é–º–µ]
        ‚úÖ [–í—Ç–æ—Ä–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ —Ä–µ–∑—é–º–µ]
        üìä **–ü—Ä–µ–¥–≤–∞—Ä–∏—Ç–µ–ª—å–Ω–∞—è –æ—Ü–µ–Ω–∫–∞:** [–ü–æ–¥—Ö–æ–¥–∏—Ç / –°—Ç–æ–∏—Ç —Ä–∞—Å—Å–º–æ—Ç—Ä–µ—Ç—å / –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –æ–ø—ã—Ç–∞]
        
        –¢–µ–∫—Å—Ç —Ä–µ–∑—é–º–µ:
        {resume_text}
        """,
        'gemini_convo_prompt': """–¢—ã –æ–ø—ã—Ç–Ω—ã–π HR-–º–µ–Ω–µ–¥–∂–µ—Ä. –ù–∏–∂–µ –ø—Ä–∏–≤–µ–¥–µ–Ω—ã –æ—Ç–≤–µ—Ç—ã –∫–∞–Ω–¥–∏–¥–∞—Ç–∞ –∏–∑ —á–∞—Ç–∞. 
        –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —ç—Ç—É –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –∏ –Ω–∞–ø–∏—à–∏ –∫—Ä–∞—Ç–∫–æ–µ –∏ —á–µ—Ç–∫–æ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ –æ –∫–∞–Ω–¥–∏–¥–∞—Ç–µ –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ.
        –ê–Ω–∞–ª–∏–∑ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –≤ —Å–ª–µ–¥—É—é—â–µ–º —Ñ–æ—Ä–º–∞—Ç–µ, –∏—Å–ø–æ–ª—å–∑—É–π —ç–º–æ–¥–∑–∏:
        ü§ñ **–û–±—â–µ–µ –∑–∞–∫–ª—é—á–µ–Ω–∏–µ:** [–ó–∞–∫–ª—é—á–µ–Ω–∏–µ –∏–∑ 2-3 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –Ω–∞ –æ—Å–Ω–æ–≤–µ –æ—Ç–≤–µ—Ç–æ–≤ –∫–∞–Ω–¥–∏–¥–∞—Ç–∞ –∏ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏—è –≤–∞–∫–∞–Ω—Å–∏–∏]
        ‚ú® **–°–∏–ª—å–Ω—ã–µ —Å—Ç–æ—Ä–æ–Ω—ã:**
        ‚úÖ [–ü–µ—Ä–≤–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ –æ—Ç–≤–µ—Ç–∞—Ö]
        ‚úÖ [–í—Ç–æ—Ä–∞—è –∫–ª—é—á–µ–≤–∞—è —Å–∏–ª—å–Ω–∞—è —Å—Ç–æ—Ä–æ–Ω–∞, –Ω–∞–π–¥–µ–Ω–Ω–∞—è –≤ –æ—Ç–≤–µ—Ç–∞—Ö]
        üìä **–ü—Ä–µ–¥–≤–∞—Ä–∏—Ç–µ–ª—å–Ω–∞—è –æ—Ü–µ–Ω–∫–∞:** [–ü–æ–¥—Ö–æ–¥–∏—Ç / –°—Ç–æ–∏—Ç —Ä–∞—Å—Å–º–æ—Ç—Ä–µ—Ç—å / –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –æ–ø—ã—Ç–∞]"""
    }
}


# --- BOTNING XOTIRASI (FSM) YANGILANDI ---
class Form(StatesGroup):
    language_selection = State()
    name = State()
    has_resume_choice = State()
    resume_upload = State()
    convo_vacancy = State()
    convo_experience = State()
    convo_salary = State()
    convo_location = State()
    convo_skills = State()
    convo_availability = State()
    convo_contact = State()

# --- ASOSIY BOT QISMI ---
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

async def get_user_lang(state: FSMContext):
    user_data = await state.get_data()
    return user_data.get('language', 'uz')

# --- BOT SUHBATLOGIKASI (QAYTA QURILDI) ---
@dp.message(Command("start"))
async def send_welcome(message: types.Message, state: FSMContext):
    await state.clear()
    language_keyboard = InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text="üá∫üáø O'zbekcha", callback_data="lang_uz")], [InlineKeyboardButton(text="üá∑üá∫ –†—É—Å—Å–∫–∏–π", callback_data="lang_ru")]])
    await message.reply(f"{texts['uz']['welcome']}\n{texts['ru']['welcome']}", reply_markup=language_keyboard)
    await state.set_state(Form.language_selection)

@dp.callback_query(Form.language_selection, F.data.startswith('lang_'))
async def process_language_selection(callback: types.CallbackQuery, state: FSMContext):
    lang = callback.data.split('_')[1]
    await state.update_data(language=lang)
    await callback.message.delete_reply_markup()
    await callback.message.answer(texts[lang]['ask_name'])
    await state.set_state(Form.name)
    await callback.answer()

@dp.message(Form.name)
async def process_name(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(name=message.text)
    
    resume_choice_keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=texts[lang]['button_yes_resume'], callback_data="has_resume_yes")],
        [InlineKeyboardButton(text=texts[lang]['button_no_resume'], callback_data="has_resume_no")]
    ])
    await message.answer(texts[lang]['ask_has_resume'], reply_markup=resume_choice_keyboard)
    await state.set_state(Form.has_resume_choice)

@dp.callback_query(Form.has_resume_choice, F.data.startswith('has_resume_'))
async def process_has_resume_choice(callback: types.CallbackQuery, state: FSMContext):
    lang = await get_user_lang(state)
    choice = callback.data.split('_')[2]
    await callback.message.delete_reply_markup()
    
    if choice == "yes":
        await callback.message.answer(texts[lang]['prompt_for_resume'])
        await state.set_state(Form.resume_upload)
    elif choice == "no":
        await callback.message.answer(texts[lang]['start_convo_application'])
        await callback.message.answer(texts[lang]['ask_vacancy'])
        await state.set_state(Form.convo_vacancy)
    await callback.answer()

# === 1-YO'L: Rezyume yuklash uchun handler ===
@dp.message(Form.resume_upload, F.document)
async def process_resume_file(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    file_mime_type = message.document.mime_type
    
    if file_mime_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        await message.reply(texts[lang]['file_error'])
        return

    await message.answer(texts[lang]['analyzing'])
    
    file_id = message.document.file_id
    file_info = await bot.get_file(file_id)
    file_bytes_io = await bot.download_file(file_info.file_path)
    
    user_data = await state.get_data()
    gemini_summary = ""

    try:
        if file_mime_type == "application/pdf":
            pdf_part = {"mime_type": "application/pdf", "data": file_bytes_io.read()}
            prompt = texts[lang]['gemini_file_prompt']
            response = await model.generate_content_async([prompt, pdf_part])
            gemini_summary = response.text
        elif file_mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document = docx.Document(file_bytes_io)
            resume_text_parts = []
            for para in document.paragraphs: resume_text_parts.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells: resume_text_parts.append(cell.text)
            resume_text = "\n".join(resume_text_parts)
            
            if not resume_text.strip():
                 gemini_summary = "DOCX faylidan matn topilmadi."
            else:
                prompt = texts[lang]['gemini_text_prompt'].format(resume_text=resume_text)
                response = await model.generate_content_async(prompt)
                gemini_summary = response.text
    except Exception as e:
        logging.error(f"Faylni tahlil qilishdagi xato: {e}")
        gemini_summary = "Faylni tahlil qilishda xatolik yuz berdi."

    hr_notification_template = texts[lang]['hr_notification_file']
    hr_summary_text = hr_notification_template.format(name=user_data.get('name'), summary=gemini_summary)
    
    if HR_GROUP_ID:
        await bot.send_message(HR_GROUP_ID, hr_summary_text, parse_mode="Markdown")
        await bot.send_document(HR_GROUP_ID, file_id)
    
    await message.answer(texts[lang]['goodbye_user'])
    await state.clear()


# === 2-YO'L: Suhbat orqali ma'lumot olish ===
@dp.message(Form.convo_vacancy)
async def process_convo_vacancy(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(vacancy=message.text)
    await message.answer(texts[lang]['ask_experience'])
    await state.set_state(Form.convo_experience)

@dp.message(Form.convo_experience)
async def process_convo_experience(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(experience=message.text)
    await message.answer(texts[lang]['ask_salary'])
    await state.set_state(Form.convo_salary)

@dp.message(Form.convo_salary)
async def process_convo_salary(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(salary=message.text)
    await message.answer(texts[lang]['ask_location'])
    await state.set_state(Form.convo_location)

@dp.message(Form.convo_location)
async def process_convo_location(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(location=message.text)
    await message.answer(texts[lang]['ask_skills'])
    await state.set_state(Form.convo_skills)

@dp.message(Form.convo_skills)
async def process_convo_skills(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    await state.update_data(skills=message.text)
    availability_keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=texts[lang]['button_yes'], callback_data="availability_yes")],
        [InlineKeyboardButton(text=texts[lang]['button_no'], callback_data="availability_no")]
    ])
    await message.answer(texts[lang]['ask_availability'], reply_markup=availability_keyboard)
    await state.set_state(Form.convo_availability)

@dp.callback_query(Form.convo_availability, F.data.startswith('availability_'))
async def process_convo_availability(callback: types.CallbackQuery, state: FSMContext):
    lang = await get_user_lang(state)
    choice = callback.data.split('_')[1]
    availability_text = texts[lang]['button_yes'] if choice == "yes" else texts[lang]['button_no']
    
    await state.update_data(availability=availability_text)
    await callback.message.delete_reply_markup()
    
    # Maxsus klaviatura yaratish
    contact_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text=texts[lang]['button_share_contact'], request_contact=True)]],
        resize_keyboard=True,
        one_time_keyboard=True
    )
    
    await callback.message.answer(texts[lang]['ask_contact'], reply_markup=contact_keyboard)
    await state.set_state(Form.convo_contact)
    await callback.answer()

# KONTAKTNI QABUL QILADIGAN funksiya
@dp.message(Form.convo_contact, F.contact)
async def process_convo_contact(message: types.Message, state: FSMContext):
    lang = await get_user_lang(state)
    contact_number = message.contact.phone_number
    await state.update_data(contact=contact_number)
    
    # Maxsus klaviaturani olib tashlash
    await message.answer(f"Raqamingiz qabul qilindi: {contact_number}", reply_markup=ReplyKeyboardRemove())
    
    await message.answer(texts[lang]['analyzing'])

    user_data = await state.get_data()
    
    # ... (Bu yerdagi Gemini tahlili va HR guruhiga yuborish logikasi o'zgarishsiz qoladi) ...
    
    candidate_summary_text = (
        f"Vakansiya: {user_data.get('vacancy')}\n"
        f"Tajribasi: {user_data.get('experience')}\n"
        f"Maosh kutilmasi: {user_data.get('salary')}\n"
        f"Manzili: {user_data.get('location')}\n"
        f"Ko'nikmalari: {user_data.get('skills')}\n"
        f"Ishga tayyorligi: {user_data.get('availability')}\n"
        f"Aloqa: {user_data.get('contact')}"
    )
    
    prompt = texts[lang]['gemini_convo_prompt']
    full_prompt = f"{prompt}\n\nNomzod javoblari:\n{candidate_summary_text}"
    
    response = await model.generate_content_async(full_prompt)
    gemini_summary = response.text
    
    hr_notification_template = texts[lang]['hr_notification_convo']
    hr_summary_text = hr_notification_template.format(**user_data, summary=gemini_summary)

    if HR_GROUP_ID:
        await bot.send_message(HR_GROUP_ID, hr_summary_text, parse_mode="Markdown")

    await message.answer(texts[lang]['goodbye_user'])
    await state.clear()

async def main():
    if not BOT_TOKEN or not GEMINI_API_KEY:
        logging.critical("Bot tokeni yoki Gemini API kaliti topilmadi!")
        return
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())